"""
Resume text extraction service.

Converts an uploaded PDF or DOCX file into clean plain text.
Heavy libraries (fitz / python-docx) are imported at call time so that
importing this module never requires them to be installed unless a
resume is actually being processed.
"""

import io


class ResumeExtractionError(Exception):
    """Raised when a resume file cannot be parsed or contains no text."""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF file's bytes using PyMuPDF."""
    try:
        import pymupdf  # PyMuPDF (the `fitz` import name is deprecated)
    except ImportError as exc:
        raise ResumeExtractionError(
            "PDF processing library (PyMuPDF) is not installed."
        ) from exc

    try:
        text_parts = []
        with pymupdf.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                text_parts.append(page.get_text())
        return "\n".join(text_parts).strip()
    except Exception as exc:
        raise ResumeExtractionError(f"Failed to read PDF file: {exc}") from exc


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX file's bytes using python-docx."""
    try:
        import docx
    except ImportError as exc:
        raise ResumeExtractionError(
            "DOCX processing library (python-docx) is not installed."
        ) from exc

    try:
        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        # Also pull text out of any tables in the document.
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        return "\n".join(paragraphs).strip()
    except Exception as exc:
        raise ResumeExtractionError(f"Failed to read DOCX file: {exc}") from exc


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    """
    Dispatch to the correct extractor based on file extension, then
    validate that meaningful text was actually extracted.
    """
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise ResumeExtractionError("Unsupported file type. Please upload a PDF or DOCX file.")

    if not text or len(text.strip()) < 20:
        raise ResumeExtractionError(
            "Could not extract meaningful text from this file. "
            "It may be empty, scanned as an image, or corrupted."
        )

    return text
